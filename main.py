from docx import Document
import xlrd
import re
if __name__ == '__main__':
    #Excel文件路径
    filename="D:/2020学术年会/2020年研究生学术年会投稿情况汇总表 - 副本.xls";
    #Word文件路径
    wordPath="C:/Users/12420/Desktop/123.docx";
    data = xlrd.open_workbook(filename);
    document = Document(wordPath);
    #选择Excel中的第一张表
    table = data.sheets()[0];
    #获得有效行数
    nrows=table.nrows;
    i=1
    while(i<nrows):
        #获得单元格内容
        paperName=table.cell_value(i,7);
        paperName=re.sub("^.*\+.*\+.*\+.*\+", "", paperName, count=0, flags=0);
        paperName = re.sub("^.*-.*-.*-.*-", "", paperName, count=0, flags=0);
        paperName = re.sub("^.*_.*_.*_.*_", "", paperName, count=0, flags=0);
        paperAuthor = table.cell_value(i, 8);
        paperKeyWordsCN = table.cell_value(i, 9);
        paperAbstractCN = table.cell_value(i, 10);
        paperKeyWordsENG = table.cell_value(i, 11);
        paperAbstractENG = table.cell_value(i, 12);
        if(paperName!=None):
            #style就是word文档样式库中的样式，注意样式名要用英文
            #添加段
            document.add_paragraph(paperName,style='nameCN');
        if (paperAuthor != None):
            document.add_paragraph(paperAuthor,style='authorCN');
        if (paperAbstractCN != None):
            #添加段和行
            document.add_paragraph("摘  要  ", style='abAndKeyCN').add_run(paperAbstractCN);
        if (paperKeyWordsCN != None):
            document.add_paragraph("关键字  ", style='abAndKeyCN').add_run(paperKeyWordsCN);
        if (paperAbstractENG != None):
            try:
                document.add_paragraph("Abstract  ", style='abAndKeyENG').add_run(paperAbstractENG);
            except ValueError as e:
                print(paperName);
                print(paperAbstractENG);
        if (paperKeyWordsENG != None):
            document.add_paragraph("Keywords  ", style='abAndKeyENG').add_run(paperKeyWordsENG);
        document.save(wordPath);# 保存文档
        i+=1;
