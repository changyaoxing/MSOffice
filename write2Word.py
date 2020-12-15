from docx import Document
import xlrd
if __name__ == '__main__':
    #Excel文件路径
    filename="D:/2020学术年会/order.xls";
    #Word文件路径
    wordPath="C:/Users/12420/Desktop/order.docx";
    data = xlrd.open_workbook(filename);
    document = Document(wordPath);
    #选择Excel中的第一张表
    table = data.sheets()[0];
    #获得有效行数
    nrows=table.nrows;
    i=0
    while(i<nrows):
        #获得单元格内容
        paperNameCN=table.cell_value(i,0);
        paperAuthorCN = table.cell_value(i, 1);
        paperKeyWordsCN = table.cell_value(i, 4);
        paperAbstractCN = table.cell_value(i, 5);
        paperNameENG = table.cell_value(i, 2);
        paperAuthorENG = table.cell_value(i, 3);
        paperKeyWordsENG = table.cell_value(i, 6);
        paperAbstractENG = table.cell_value(i, 7);
        if(paperNameCN!="NaN"):
            #style就是word文档样式库中的样式，注意样式名要用英文
            #添加段
            document.add_paragraph(paperNameCN,style='nameCN');
        if (paperAuthorCN != "NaN"):
            document.add_paragraph(paperAuthorCN,style='authorCN');
        if (paperAbstractCN != "NaN"):
            #添加段和行
            document.add_paragraph("摘  要  ", style='abAndKeyCN').add_run(paperAbstractCN);
        if (paperKeyWordsCN != "NaN"):
            document.add_paragraph("关键字  ", style='abAndKeyCN').add_run(paperKeyWordsCN);

        if (paperNameENG != "NaN"):
            # style就是word文档样式库中的样式，注意样式名要用英文
            # 添加段
            document.add_paragraph(paperNameENG, style='nameENG');
        if (paperAuthorENG != "NaN"):
            document.add_paragraph(paperAuthorENG, style='authorENG');
        if (paperAbstractENG != "NaN"):
            try:
                document.add_paragraph("Abstract  ", style='abAndKeyENG').add_run(paperAbstractENG);
            except ValueError as e:
                print(paperNameCN);
                print(paperAbstractENG);
        if (paperKeyWordsENG != "NaN"):
            document.add_paragraph("Keywords  ", style='abAndKeyENG').add_run(paperKeyWordsENG);
        document.save(wordPath);# 保存文档
        i+=1;