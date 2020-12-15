import xlrd
import os
from docx import Document
import xlwt
import re
if __name__ == '__main__':
    excelName = "D:/2020学术年会/paperName.xls";
    dirPath="D:/2020学术年会/2020年计算机学院研究生学术年会论文";
    data = xlrd.open_workbook(excelName);
    table = data.sheets()[0];
    # 获得有效行数
    nrows = table.nrows;
    i = 0;
    excelNameList=[];
    while (i < nrows):
        # 获得单元格内容
        paperName = table.cell_value(i, 0);
        excelNameList.append(paperName);
        i+=1;
    fileNameList=[];
    filePathList=[];
    paperNameDic={};
    paperAuthorDic={};
    for root, dirs, files in os.walk(dirPath):
        for fileName in files:
            if(re.search("doc$",fileName)!=None or re.search("docx$",fileName)!=None):

                #     fileName=fileName.replace("doc","docx")
                filePath=os.path.join(root, fileName);
                filePath=filePath.replace("\\", "/")
                # if (re.search("doc$", fileName) != None):
                #     os.rename(filePath,filePath.replace("doc","docx"))
                #     filePath=filePath.replace("doc", "docx")
                filePathList.append(filePath);
                # print(filePath)
                fileName = re.sub("^.*\+.*\+.*\+.*\+", "", fileName, count=0, flags=0);
                fileName = re.sub("^.*-.*-.*-.*-", "", fileName, count=0, flags=0);
                fileName = re.sub("^.*_.*_.*_.*_", "", fileName, count=0, flags=0);
                fileNameList.append(fileName)
                print(fileName)
                try:
                    document = Document(filePath);
                    for i in range(len(document.paragraphs)):
                        para=document.paragraphs[i];
                        flag=True;

                        t=para.text;
                        if i+1<len(document.paragraphs):
                            a=document.paragraphs[i+1].text;
                        else:
                            a="";

                        if(t=="" or t=="\n"):
                            flag = False;
                        else:
                            for ch in t:
                                if u'\u4e00' <= ch <= u'\u9fff':
                                    flag=False;
                                    break;
                        if flag:
                            paperNameDic[fileName]=t;
                            paperAuthorDic[fileName]=a;
                            print(t+"\r\n ###############################");
                            break;
                    try:
                        document.save(filePath);
                    except ValueError as e:
                        print("---------",e, filePath)
                except ValueError as e:
                    print("~~~~~~~~~~~~~~~~",e,filePath)
                # paperNameDic[fileName]=


    workbook = xlwt.Workbook();
    worksheet = workbook.add_sheet('My Sheet')
    i=0;
    for x,y in paperNameDic.items():
        worksheet.write(i, 1, x);
        worksheet.write(i, 2, y);
        worksheet.write(i,3,paperAuthorDic[x]);
        workbook.save(excelName);
        i+=1;
    print(paperNameDic);

